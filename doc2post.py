# -*- coding: utf-8 -*-
"""
Word 文档转博客文章
读取 posts/ 目录下的 .docx 文件，解析元信息和正文，
转换为干净的 HTML 并插入 js/main.js 的 posts 数组。

Word 文档格式要求：
- 第1行：标题（必填）
- 第2行：摘要（可留空）
- 第3行：分类（技术/生活/读书，可留空默认技术）
- 第4行：日期（YYYY-MM-DD，可留空默认今天）
- 第5行起：正文

正文样式约定：
- "标题1"~"标题3" 或 Heading 1~3 → <h2>/<h3>
- 等宽字体（Consolas/Courier New/仿宋_GB2312 等等宽字体）→ <code> 或 <pre><code>
- 图片自动提取到 images/posts/ 目录
"""
import os
import re
import sys
import shutil
import datetime
from io import BytesIO

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print('[错误] 需要安装 python-docx 库')
    print('运行: pip install python-docx')
    sys.exit(1)


ROOT = os.path.dirname(os.path.abspath(__file__))
JS_PATH = os.path.join(ROOT, 'js', 'main.js')
POSTS_DIR = os.path.join(ROOT, 'posts')
IMAGES_DIR = os.path.join(ROOT, 'images', 'posts')

TAG_MAP = {
    '技术': ('tech', '技术'),
    '生活': ('life', '生活'),
    '读书': ('reading', '读书'),
    'tech': ('tech', '技术'),
    'life': ('life', '生活'),
    'reading': ('reading', '读书'),
}

MONOSPACE_FONTS = {
    'consolas', 'courier new', 'courier', 'menlo', 'monaco',
    'source code pro', 'fira code', 'jetbrains mono',
    '仿宋_gb2312', '仿宋', '等线', 'dengxian',
}


def escape_html(text):
    """HTML 特殊字符转义"""
    return (text.replace('&', '&amp;')
                .replace('<', '&lt;')
                .replace('>', '&gt;')
                .replace('"', '&quot;'))


def escape_js_string(text):
    """JS 字符串转义（用于 title/excerpt 等字段）"""
    return (text.replace('\\', '\\\\')
                .replace('"', '\\"')
                .replace('\r', '')
                .replace('\n', ' ')
                .strip())


def is_monospace(run):
    """判断 run 是否使用等宽字体"""
    if not run.font.name:
        return False
    font_name = run.font.name.lower()
    # 检查西文字体
    if font_name in MONOSPACE_FONTS:
        return True
    # 检查东亚字体
    rpr = run._element.find(qn('w:rPr'))
    if rpr is not None:
        rfonts = rpr.find(qn('w:rFonts'))
        if rfonts is not None:
            east_asia = rfonts.get(qn('w:eastAsia'), '').lower()
            if east_asia in MONOSPACE_FONTS:
                return True
    return False


def extract_images(doc, post_id):
    """提取文档中的所有图片，返回 {relationship_id: filename} 映射"""
    image_map = {}
    os.makedirs(IMAGES_DIR, exist_ok=True)

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            ext = os.path.splitext(rel.target_ref)[1] or '.png'
            filename = f'post{post_id}_{rel_id}{ext}'
            filepath = os.path.join(IMAGES_DIR, filename)
            with open(filepath, 'wb') as f:
                f.write(image_data)
            image_map[rel_id] = f'images/posts/{filename}'

    return image_map


def parse_metadata(paragraphs):
    """从文档前4行解析元信息"""
    meta = {
        'title': '',
        'excerpt': '',
        'tag': 'tech',
        'tag_label': '技术',
        'date': datetime.date.today().isoformat(),
    }

    lines = []
    for p in paragraphs[:4]:
        text = p.text.strip()
        lines.append(text)

    if len(lines) >= 1 and lines[0]:
        meta['title'] = lines[0]
    if len(lines) >= 2 and lines[1]:
        meta['excerpt'] = lines[1]
    if len(lines) >= 3 and lines[2]:
        tag_input = lines[2].strip()
        if tag_input in TAG_MAP:
            meta['tag'], meta['tag_label'] = TAG_MAP[tag_input]
    if len(lines) >= 4 and lines[3]:
        date_str = lines[3].strip()
        try:
            datetime.date.fromisoformat(date_str)
            meta['date'] = date_str
        except ValueError:
            pass

    return meta


def convert_paragraph(para, image_map):
    """将单个段落转换为 HTML"""
    style_name = para.style.name.lower() if para.style else ''

    # 标题样式
    if 'heading 1' in style_name or '标题 1' in style_name:
        text = escape_html(para.text.strip())
        if text:
            return f'<h2>{text}</h2>'
    elif 'heading 2' in style_name or '标题 2' in style_name:
        text = escape_html(para.text.strip())
        if text:
            return f'<h3>{text}</h3>'
    elif 'heading 3' in style_name or '标题 3' in style_name:
        text = escape_html(para.text.strip())
        if text:
            return f'<h4>{text}</h4>'

    # 检查是否包含图片
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        html_parts = []
        for drawing in drawings:
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is not None:
                embed_id = blip.get(qn('r:embed'))
                if embed_id and embed_id in image_map:
                    img_src = image_map[embed_id]
                    html_parts.append(f'<img src="{img_src}" alt="文章配图">')
        if html_parts:
            return ''.join(html_parts)

    # 检查是否是代码块（整段等宽字体）
    runs = para.runs
    if runs and all(is_monospace(r) for r in runs if r.text.strip()):
        code_text = escape_html(para.text)
        return f'<pre><code>{code_text}</code></pre>'

    # 普通段落：处理行内代码和文本
    if not para.text.strip():
        return ''

    html_parts = []
    for run in runs:
        text = escape_html(run.text)
        if not text:
            continue
        if is_monospace(run):
            html_parts.append(f'<code>{text}</code>')
        elif run.bold:
            html_parts.append(f'<strong>{text}</strong>')
        elif run.italic:
            html_parts.append(f'<em>{text}</em>')
        else:
            html_parts.append(text)

    content = ''.join(html_parts)
    if content.strip():
        return f'<p>{content}</p>'
    return ''


def convert_docx_to_html(docx_path):
    """将 Word 文档转换为 HTML，返回 (metadata, html_content)"""
    doc = Document(docx_path)

    # 获取最大 id
    try:
        with open(JS_PATH, 'rb') as f:
            js_text = f.read().decode('utf-8')
        ids = [int(x) for x in re.findall(r'(?m)^\s*id:\s*(\d+)\s*,', js_text)]
        new_id = max(ids) + 1 if ids else 1
    except FileNotFoundError:
        print('[错误] 找不到 js/main.js')
        return None, None

    # 提取图片
    image_map = extract_images(doc, new_id)

    # 解析元信息
    meta = parse_metadata(doc.paragraphs)
    if not meta['title']:
        print('[错误] 文档第1行必须写标题')
        return None, None

    meta['id'] = new_id

    # 转换正文（跳过前4行元信息）
    html_parts = []
    for para in doc.paragraphs[4:]:
        html = convert_paragraph(para, image_map)
        if html:
            html_parts.append(html)

    html_content = '\n            '.join(html_parts)
    return meta, html_content


def insert_to_mainjs(meta, html_content):
    """将文章数据插入 main.js"""
    with open(JS_PATH, 'rb') as f:
        raw = f.read()
    text = raw.decode('utf-8')

    marker = 'const posts = ['
    pos = text.find(marker)
    if pos == -1:
        print('[错误] 在 js/main.js 里找不到 posts 数组')
        return False

    nl = '\r\n' if raw.count(b'\r\n') * 2 > raw.count(b'\n') else '\n'

    lines = [
        '    {',
        '        id: %d,' % meta['id'],
        '        title: "%s",' % escape_js_string(meta['title']),
        '        excerpt: "%s",' % escape_js_string(meta['excerpt'] or '暂无摘要'),
        '        date: "%s",' % meta['date'],
        '        tag: "%s",' % meta['tag'],
        '        tagLabel: "%s",' % meta['tag_label'],
        '        content: `',
        '            %s' % html_content.replace('\n', '\n            '),
        '        `',
        '    },',
    ]
    block = nl.join(lines) + nl

    insert_at = text.find('\n', pos) + 1
    new_text = text[:insert_at] + block + text[insert_at:]

    with open(JS_PATH, 'w', encoding='utf-8', newline='') as f:
        f.write(new_text)

    # 自检
    with open(JS_PATH, 'rb') as f:
        check = f.read().decode('utf-8')
    new_ids = re.findall(r'(?m)^\s*id:\s*(\d+)\s*,', check)
    if str(meta['id']) not in new_ids:
        print('[警告] 写入结果异常，请手动检查 js/main.js')
        return False

    return True


def main():
    os.makedirs(POSTS_DIR, exist_ok=True)

    # 列出可用的 docx 文件
    docx_files = [f for f in os.listdir(POSTS_DIR) if f.endswith('.docx')]
    if not docx_files:
        print('=' * 44)
        print('  Word 转博客文章')
        print('=' * 44)
        print()
        print('posts/ 目录下没有找到 .docx 文件。')
        print()
        print('使用方法：')
        print('  1. 在 posts/ 目录下新建 Word 文档')
        print('  2. 按以下格式填写：')
        print('     第1行：文章标题（必填）')
        print('     第2行：摘要（可留空）')
        print('     第3行：分类（技术/生活/读书，可留空默认技术）')
        print('     第4行：日期（YYYY-MM-DD，可留空默认今天）')
        print('     第5行起：正文')
        print('  3. 再次运行本脚本')
        return 0

    print('=' * 44)
    print('  Word 转博客文章')
    print('=' * 44)
    print()
    print('找到以下文档：')
    for i, f in enumerate(docx_files, 1):
        print(f'  [{i}] {f}')
    print()

    choice = input('选择要转换的文件编号（直接回车转换全部）: ').strip()

    if choice:
        try:
            idx = int(choice) - 1
            if idx < 0 or idx >= len(docx_files):
                print('[错误] 无效的编号')
                return 1
            selected = [docx_files[idx]]
        except ValueError:
            print('[错误] 请输入数字')
            return 1
    else:
        selected = docx_files

    success_count = 0
    for filename in selected:
        filepath = os.path.join(POSTS_DIR, filename)
        print(f'\n正在转换: {filename}')

        meta, html = convert_docx_to_html(filepath)
        if meta is None:
            continue

        if insert_to_mainjs(meta, html):
            print(f'  ✓ 已插入: ID={meta["id"]}, 标题={meta["title"]}')
            success_count += 1

            # 转换成功后移动原文件到 processed/
            processed_dir = os.path.join(POSTS_DIR, 'processed')
            os.makedirs(processed_dir, exist_ok=True)
            dest = os.path.join(processed_dir, filename)
            shutil.move(filepath, dest)
            print(f'  ✓ 原文件已移至 posts/processed/')

    print()
    print('=' * 44)
    print(f'  完成！成功转换 {success_count} 篇文章')
    print('=' * 44)
    print()
    print('接下来：')
    print('  1. 本地刷新 posts.html 预览效果')
    print('  2. 双击 deploy.bat 发布上线')

    return 0


if __name__ == '__main__':
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print('\n已取消。')
        sys.exit(1)
